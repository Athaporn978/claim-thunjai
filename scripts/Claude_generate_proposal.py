import os
from docx import Document
from docx.shared import Inches, Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_TAB_ALIGNMENT, WD_TAB_LEADER
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import nsdecls, qn

# ----------------------------------------------------------------------------
# Design tokens
# ----------------------------------------------------------------------------
FONT = "TH Sarabun New"
BLUE = RGBColor(0, 113, 227)        # #0071e3 — primary accent
NAVY = RGBColor(10, 37, 64)         # #0a2540 — deep banner navy
NAVY2 = RGBColor(15, 76, 129)       # #0f4c81 — secondary navy
TEXT = RGBColor(45, 52, 64)         # #2d3440 — body text
MUTED = RGBColor(107, 122, 143)     # #6b7a8f — muted grey-blue
WHITE = RGBColor(255, 255, 255)
LIGHT_TINT = "EAF3FE"               # very light blue row tint
BORDER_GREY = "D7E2EF"

PAGE_NUM = {}  # heading text -> page number (static, pre-computed / hand-tracked)


def create_proposal_docx():
    doc = Document()

    normal = doc.styles["Normal"]
    normal.font.name = FONT
    normal.font.size = Pt(16)
    normal.font.color.rgb = TEXT
    rpr = normal.element.get_or_add_rPr()
    rFonts = rpr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts')
        rpr.append(rFonts)
    for attr in ('w:ascii', 'w:hAnsi', 'w:cs', 'w:eastAsia'):
        rFonts.set(qn(attr), FONT)

    for section in doc.sections:
        section.top_margin = Cm(1.6)
        section.bottom_margin = Cm(1.8)
        section.left_margin = Cm(2.2)
        section.right_margin = Cm(2.2)

    # ---- low level helpers --------------------------------------------
    def set_font(run, size=16, bold=False, italic=False, color=TEXT):
        run.font.name = FONT
        run.font.size = Pt(size)
        run.font.bold = bold
        run.font.italic = italic
        run.font.color.rgb = color
        rPr = run._r.get_or_add_rPr()
        rFonts = OxmlElement('w:rFonts')
        for attr in ('w:ascii', 'w:hAnsi', 'w:cs', 'w:eastAsia'):
            rFonts.set(qn(attr), FONT)
        rPr.append(rFonts)

    def para(align=WD_ALIGN_PARAGRAPH.LEFT, before=0, after=6, line=1.2):
        p = doc.add_paragraph()
        p.alignment = align
        p.paragraph_format.space_before = Pt(before)
        p.paragraph_format.space_after = Pt(after)
        p.paragraph_format.line_spacing = line
        return p

    def shade(cell, hex_color):
        cell._tc.get_or_add_tcPr().append(parse_xml(f'<w:shd {nsdecls("w")} w:fill="{hex_color}"/>'))

    def no_borders(table):
        tblPr = table._tbl.tblPr
        borders = OxmlElement('w:tblBorders')
        for edge in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
            el = OxmlElement(f'w:{edge}')
            el.set(qn('w:val'), 'nil')
            borders.append(el)
        tblPr.append(borders)

    def grid_borders(table, color=BORDER_GREY, sz=4):
        tblPr = table._tbl.tblPr
        borders = OxmlElement('w:tblBorders')
        for edge in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
            el = OxmlElement(f'w:{edge}')
            el.set(qn('w:val'), 'single')
            el.set(qn('w:sz'), str(sz))
            el.set(qn('w:space'), '0')
            el.set(qn('w:color'), color)
            borders.append(el)
        tblPr.append(borders)

    def cell_margins(cell, top=100, bottom=100, left=180, right=180):
        tcPr = cell._tc.get_or_add_tcPr()
        mar = OxmlElement('w:tcMar')
        for tag, val in (('top', top), ('bottom', bottom), ('left', left), ('right', right)):
            el = OxmlElement(f'w:{tag}')
            el.set(qn('w:w'), str(val))
            el.set(qn('w:type'), 'dxa')
            mar.append(el)
        tcPr.append(mar)

    def set_col_widths(table, widths):
        table.autofit = False
        for row in table.rows:
            for i, w in enumerate(widths):
                row.cells[i].width = w
        for i, w in enumerate(widths):
            table.columns[i].width = w

    # ---- content helpers -------------------------------------------------
    def chapter_banner(number, title_th, title_en):
        """Full-width dark navy banner used for top-level chapters."""
        tbl = doc.add_table(rows=1, cols=2)
        tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
        no_borders(tbl)
        set_col_widths(tbl, [Cm(2.6), Cm(12.6)])
        c1, c2 = tbl.rows[0].cells
        shade(c1, "0A2540")
        shade(c2, "0071E3")
        cell_margins(c1, top=140, bottom=140, left=160, right=80)
        cell_margins(c2, top=140, bottom=140, left=220, right=160)

        p1 = c1.paragraphs[0]
        r1 = p1.add_run(number)
        set_font(r1, size=30, bold=True, color=WHITE)
        p1.alignment = WD_ALIGN_PARAGRAPH.CENTER

        p2 = c2.paragraphs[0]
        r2 = p2.add_run(title_th)
        set_font(r2, size=21, bold=True, color=WHITE)
        p2.paragraph_format.space_after = Pt(2)
        if title_en:
            p3 = c2.add_paragraph()
            r3 = p3.add_run(title_en)
            set_font(r3, size=13, italic=True, color=RGBColor(210, 228, 250))
        doc.add_paragraph().paragraph_format.space_after = Pt(4)
        return tbl

    def sub_heading(text, key=None):
        p = para(before=12, after=6)
        pPr = p._p.get_or_add_pPr()
        pPr.append(parse_xml(
            r'<w:pBdr xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">'
            r'<w:left w:val="single" w:sz="24" w:space="6" w:color="0071E3"/></w:pBdr>'))
        pPr.append(parse_xml(f'<w:ind xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main" w:left="200"/>'))
        r = p.add_run(text)
        set_font(r, size=18, bold=True, color=NAVY2)
        if key:
            PAGE_NUM[key] = None  # filled later by track_page()
        return p

    def body(text, italic=False, after=8, align=WD_ALIGN_PARAGRAPH.LEFT):
        p = para(after=after, align=align)
        r = p.add_run(text)
        set_font(r, size=16, italic=italic)
        return p

    def bullet(text, prefix=""):
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Cm(0.9)
        p.paragraph_format.first_line_indent = Cm(-0.5)
        p.paragraph_format.space_after = Pt(5)
        p.paragraph_format.line_spacing = 1.15
        rm = p.add_run("■  ")
        set_font(rm, size=11, bold=True, color=BLUE)
        if prefix:
            rp = p.add_run(prefix)
            set_font(rp, size=16, bold=True, color=NAVY2)
        rt = p.add_run(text)
        set_font(rt, size=16)
        return p

    def callout(lines, title=""):
        tbl = doc.add_table(rows=1, cols=1)
        tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
        cell = tbl.cell(0, 0)
        shade(cell, "F2F8FF")
        cell_margins(cell, top=80, bottom=80, left=260, right=220)
        tcPr = cell._tc.get_or_add_tcPr()
        tcPr.append(parse_xml(
            r'<w:tcBorders xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">'
            r'<w:left w:val="single" w:sz="32" w:color="0071E3"/>'
            r'<w:top w:val="single" w:sz="4" w:color="D7E2EF"/>'
            r'<w:bottom w:val="single" w:sz="4" w:color="D7E2EF"/>'
            r'<w:right w:val="single" w:sz="4" w:color="D7E2EF"/></w:tcBorders>'))
        p = cell.paragraphs[0]
        if title:
            r = p.add_run("◆  " + title)
            set_font(r, size=16, bold=True, color=NAVY2)
        for i, line in enumerate(lines):
            pp = p if (i == 0 and not title) else cell.add_paragraph()
            pp.paragraph_format.space_after = Pt(2)
            pp.paragraph_format.line_spacing = 1.05
            r = pp.add_run(line)
            set_font(r, size=14, color=RGBColor(30, 58, 92))

    def data_table(headers, rows, widths, header_color="0071E3"):
        tbl = doc.add_table(rows=1 + len(rows), cols=len(headers))
        tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
        set_col_widths(tbl, widths)
        grid_borders(tbl)
        hdr = tbl.rows[0].cells
        for i, h in enumerate(headers):
            shade(hdr[i], header_color)
            cell_margins(hdr[i], top=60, bottom=60, left=140, right=140)
            p = hdr[i].paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            r = p.add_run(h)
            set_font(r, size=13, bold=True, color=WHITE)
        for ridx, row in enumerate(rows, start=1):
            cells = tbl.rows[ridx].cells
            if ridx % 2 == 0:
                for c in cells:
                    shade(c, LIGHT_TINT)
            for cidx, val in enumerate(row):
                cell_margins(cells[cidx], top=36, bottom=36, left=140, right=140)
                p = cells[cidx].paragraphs[0]
                p.paragraph_format.line_spacing = 1.0
                p.paragraph_format.space_after = Pt(0)
                r = p.add_run(val)
                set_font(r, size=12, bold=(cidx == 0), color=NAVY2 if cidx == 0 else TEXT)
        return tbl

    def track_page(key, page):
        PAGE_NUM[key] = page

    # ------------------------------------------------------------------
    # PAGE 1 — COVER
    # ------------------------------------------------------------------
    top_band = doc.add_table(rows=1, cols=1)
    no_borders(top_band)
    c = top_band.cell(0, 0)
    shade(c, "0A2540")
    cell_margins(c, top=200, bottom=200, left=260, right=260)
    p = c.paragraphs[0]
    r = p.add_run("บริษัท เทคทันใจ อินโนเวชั่น จำกัด   |   TECHTHUNJAI INNOVATION CO., LTD.")
    set_font(r, size=13, bold=True, color=WHITE)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    for _ in range(2):
        para(after=6)

    p = para(align=WD_ALIGN_PARAGRAPH.CENTER, after=6)
    r = p.add_run("EXECUTIVE  PROJECT  PROPOSAL")
    set_font(r, size=14, bold=True, color=BLUE)

    p = para(align=WD_ALIGN_PARAGRAPH.CENTER, after=10)
    r = p.add_run("ข้อเสนอโครงการระบบบริหารจัดการเคลม\nและควบคุมราคาซ่อมด้วยปัญญาประดิษฐ์")
    set_font(r, size=28, bold=True, color=NAVY)

    p = para(align=WD_ALIGN_PARAGRAPH.CENTER, after=6)
    r = p.add_run("Claim-Thunjai")
    set_font(r, size=19, bold=True, color=BLUE)

    p = para(align=WD_ALIGN_PARAGRAPH.CENTER, after=16)
    r = p.add_run("AI-Powered Repair Quotation & Price Control Platform")
    set_font(r, size=14, italic=True, color=MUTED)

    divider = para(align=WD_ALIGN_PARAGRAPH.CENTER, after=16)
    dPr = divider._p.get_or_add_pPr()
    dPr.append(parse_xml(
        r'<w:pBdr xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">'
        r'<w:bottom w:val="single" w:sz="18" w:space="1" w:color="0071E3"/></w:pBdr>'))

    info_tbl = doc.add_table(rows=4, cols=2)
    no_borders(info_tbl)
    set_col_widths(info_tbl, [Cm(4.2), Cm(9.6)])
    info_rows = [
        ("นำเสนอสำหรับ", "คณะผู้บริหารระดับสูง H Technology CO., LTD. (CEO & Board of Directors)"),
        ("จัดทำโดย", "บริษัท เทคทันใจ อินโนเวชั่น จำกัด (TECHTHUNJAI INNOVATION CO., LTD.)"),
        ("ช่องทางติดต่อ", "Email: athaporn@techthunjai.com   |   Line OA: @techthunjai   |   โทร: 065-882-8333"),
        ("วันที่นำเสนอ", "4 สิงหาคม 2569 (4 August 2026)"),
    ]
    for i, (label, val) in enumerate(info_rows):
        shade(info_tbl.rows[i].cells[0], "F2F8FF")
        shade(info_tbl.rows[i].cells[1], "F2F8FF")
        cell_margins(info_tbl.rows[i].cells[0], top=100, bottom=100, left=200, right=140)
        cell_margins(info_tbl.rows[i].cells[1], top=100, bottom=100, left=140, right=200)
        p0 = info_tbl.rows[i].cells[0].paragraphs[0]
        r0 = p0.add_run(label)
        set_font(r0, size=14, bold=True, color=NAVY2)
        p1 = info_tbl.rows[i].cells[1].paragraphs[0]
        r1 = p1.add_run(val)
        set_font(r1, size=14, color=TEXT)

    doc.add_page_break()

    # ------------------------------------------------------------------
    # PAGE 2 — TABLE OF CONTENTS (static, pre-computed page numbers below)
    # ------------------------------------------------------------------
    p = para(before=4, after=18, align=WD_ALIGN_PARAGRAPH.LEFT)
    r = p.add_run("สารบัญ")
    set_font(r, size=26, bold=True, color=BLUE)
    r2 = p.add_run("   Table of Contents")
    set_font(r2, size=14, italic=True, color=MUTED)

    toc_entries = [
        ("บทสรุปผู้บริหาร (Executive Summary)", 3, True),
        ("หมวดที่ 1   โมดูลและฟีเจอร์ทั้งหมดของระบบ (System Modules & Features)", 4, True),
        ("1.1   ระบบยืนยันตัวตนและบริหารสิทธิ์การเข้าถึงตามสาขา (Authentication & Branch RBAC)", 4, False),
        ("1.2   ระบบสร้างและอนุมัติใบเคลม/ใบเสนอราคา (Quotation Lifecycle & Approval Workflow)", 4, False),
        ("1.3   โมดูล AI สกัดข้อมูลจากเอกสารใบเสนอราคา (AI Document Extraction)", 5, False),
        ("1.4   โมดูล AI วิเคราะห์ภาพความเสียหายของรถยนต์ (AI Damage Photo Analysis)", 5, False),
        ("1.5   ระบบตรวจสภาพรถด้วยตนเองผ่านลิงก์ (Customer Self-Inspection Portal)", 6, False),
        ("1.6   แคตตาล็อกราคากลางและการนำเข้าข้อมูล (Price Catalog & Excel Import)", 6, False),
        ("1.7   ระบบบริหารจัดการส่วนกลาง (Admin: Branches / Roles / Users / Workflow)", 7, False),
        ("1.8   รายงานและสถิติเชิงบริหาร (Reports & Analytics)", 7, False),
        ("1.9   ระบบบันทึกประวัติการทำงาน (Audit Trail & Logging)", 7, False),
        ("หมวดที่ 2   สถาปัตยกรรมระบบและ API ที่ใช้งาน (System Architecture & APIs)", 8, True),
        ("หมวดที่ 3   เทคโนโลยีทั้งหมดที่ใช้ในโปรเจกต์ (Technology Stack)", 9, True),
        ("หมวดที่ 4   ข้อเสนอราคา — Pay-Per-Completed-Transaction Model", 10, True),
    ]

    for text, page, is_chapter in toc_entries:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(7 if is_chapter else 4)
        p.paragraph_format.space_before = Pt(10 if is_chapter else 0)
        tab_stops = p.paragraph_format.tab_stops
        tab_stops.add_tab_stop(Cm(15.8), WD_TAB_ALIGNMENT.RIGHT, WD_TAB_LEADER.DOTS)
        r1 = p.add_run(text + "\t")
        set_font(r1, size=16 if is_chapter else 15, bold=is_chapter, color=NAVY2 if is_chapter else TEXT)
        r2 = p.add_run(str(page))
        set_font(r2, size=16 if is_chapter else 15, bold=True, color=BLUE)

    doc.add_page_break()

    # ------------------------------------------------------------------
    # PAGE 3 — EXECUTIVE SUMMARY
    # ------------------------------------------------------------------
    chapter_banner("00", "บทสรุปผู้บริหาร", "Executive Summary")
    body("Claim-Thunjai คือแพลตฟอร์มบริหารจัดการเคลมประกันภัยรถยนต์และควบคุมราคาซ่อม ที่นำปัญญาประดิษฐ์ (Claude AI) "
         "เข้ามาช่วยอ่านและถอดข้อมูลจากใบเสนอราคาซ่อมโดยอัตโนมัติ เปรียบเทียบกับราคากลางมาตรฐาน และควบคุมกระบวนการอนุมัติ "
         "ผ่านระบบสิทธิ์การเข้าถึงตามสาขาและลำดับการอนุมัติที่ตั้งค่าได้ ช่วยลดภาระงานของเจ้าหน้าที่คุมราคา เพิ่มความรวดเร็ว "
         "และเพิ่มความโปร่งใสในการตรวจสอบทุกขั้นตอน")
    callout([
        "• ลดเวลาการอ่านและกรอกข้อมูลใบเสนอราคาจากหลักนาทีเหลือเพียงไม่กี่วินาทีต่อเอกสาร ด้วย AI Extraction",
        "• ควบคุมราคาซ่อมเทียบกับราคากลางมาตรฐานโดยอัตโนมัติ พร้อมระบบคำนวณส่วนต่าง (Saving) ทุกเคส",
        "• รองรับหลายสาขาแบบ Dynamic RBAC พร้อมลำดับการอนุมัติตามวงเงินที่ตั้งค่าได้ (Workflow Setting)",
        "• มีระบบบันทึก Audit Log ครบทุกการเปลี่ยนแปลง เพื่อการตรวจสอบย้อนหลังและความโปร่งใส",
        "• คิดค่าบริการเฉพาะเคสที่อนุมัติสำเร็จเท่านั้น (Pay-Per-Completed-Transaction) ความเสี่ยงทางธุรกิจต่ำที่สุด",
    ], title="จุดเด่นและคุณค่าทางธุรกิจหลัก")

    doc.add_page_break()

    # ------------------------------------------------------------------
    # PAGE 4 — CHAPTER 1 INTRO + 1.1 + 1.2
    # ------------------------------------------------------------------
    chapter_banner("01", "โมดูลและฟีเจอร์ทั้งหมดของระบบ", "System Modules & Features")

    sub_heading("1.1   ระบบยืนยันตัวตนและบริหารสิทธิ์การเข้าถึงตามสาขา")
    body("ระบบเข้าสู่ระบบด้วยอีเมลหรือรหัสพนักงาน ผูกสิทธิ์การใช้งานกับบทบาท (Role) และสาขา (Branch) ของผู้ใช้แต่ละคน:", after=4)
    bullet("ผู้ใช้แต่ละคนเห็นเฉพาะข้อมูลของสาขาตนเอง ยกเว้นผู้ดูแลระบบระดับสูงที่เห็นภาพรวมทุกสาขา", prefix="Dynamic Branch Scoping — ")
    bullet("กำหนดบทบาทและสิทธิ์การเข้าถึงแยกเป็นรายฟังก์ชันผ่านหน้าจัดการบทบาท", prefix="Role-Based Access Control — ")
    bullet("ปิดการใช้งานสาขาหรือบทบาทใดก็ได้ทันที ผู้ใช้ในกลุ่มนั้นจะไม่สามารถเข้าระบบได้ทันที", prefix="Instant Deactivation — ")

    sub_heading("1.2   ระบบสร้างและอนุมัติใบเคลม/ใบเสนอราคา")
    body("แกนหลักของระบบ ครอบคลุมตั้งแต่สร้างเคสจนถึงอนุมัติเสร็จสมบูรณ์:", after=4)
    bullet("แบบร่าง → รอตรวจสอบ → รอผู้มีอำนาจอนุมัติ → อนุมัติ/ปฏิเสธ พร้อมบันทึกเหตุผลทุกขั้นตอน", prefix="สถานะเคสครบวงจร — ")
    bullet("หากยอดเคสเกินวงเงินที่กำหนด ระบบส่งต่อให้ผู้มีอำนาจอนุมัติสูงขึ้นโดยอัตโนมัติตามเงื่อนไขที่ตั้งค่าได้", prefix="Threshold-Based Routing — ")
    bullet("เปรียบเทียบราคาที่อู่เสนอกับราคาที่ระบบควบคุม พร้อมคำนวณยอดประหยัดสุทธิของแต่ละเคสอัตโนมัติ", prefix="Price Comparison Engine — ")
    bullet("ล็อกไม่ให้แก้ไขเคสที่อนุมัติเสร็จสิ้นแล้ว เพื่อรักษาความถูกต้องของข้อมูลที่อนุมัติ", prefix="Approved-Case Edit Lock — ")

    doc.add_page_break()

    # ------------------------------------------------------------------
    # PAGE 5 — 1.3 + 1.4
    # ------------------------------------------------------------------
    sub_heading("1.3   โมดูล AI สกัดข้อมูลจากเอกสารใบเสนอราคา")
    body("อัปโหลดใบเสนอราคาซ่อมเป็นไฟล์ PDF หรือรูปถ่าย ระบบจะอ่านและกรอกข้อมูลลงฟอร์มให้อัตโนมัติ:", after=4)
    bullet("ถอดข้อความจาก PDF ด้วยเทคนิค text parsing ก่อนเป็นลำดับแรก เพื่อความรวดเร็วและประหยัดต้นทุน", prefix="Hybrid Extraction Strategy — ")
    bullet("หากเอกสารเป็นภาพสแกนหรือไม่มีข้อความให้ถอดโดยตรง ระบบส่งต่อให้ Claude Vision วิเคราะห์ภาพแทน", prefix="AI Vision Fallback — ")
    bullet("ชื่อ-ที่อยู่-เบอร์โทรลูกค้า, ข้อมูลกรมธรรม์/บริษัทประกัน, ข้อมูลอู่ซ่อม, รายการค่าแรงและอะไหล่พร้อมราคา", prefix="ข้อมูลที่สกัดได้ครอบคลุม — ")
    bullet("ตัวถอดรายการซ่อมแบบยืดหยุ่น รองรับรูปแบบเอกสารภาษาไทยที่หลากหลาย ไม่ตกหล่นรายการ", prefix="Flexible Item Parsing — ")

    sub_heading("1.4   โมดูล AI วิเคราะห์ภาพความเสียหายของรถยนต์")
    body("วิเคราะห์รูปถ่ายความเสียหายของรถยนต์ด้วย Claude Vision:", after=4)
    bullet("ระบุตำแหน่งชิ้นส่วนที่เสียหายพร้อมกรอบตำแหน่ง (bounding box) บนภาพ", prefix="Damage Localization — ")
    bullet("จัดระดับความรุนแรงของความเสียหายพร้อมคำอธิบายภาษาไทย/อังกฤษ และค่าความเชื่อมั่นของผลวิเคราะห์", prefix="Severity & Confidence — ")
    bullet("เชื่อมโยงผลตรวจกับแคตตาล็อกราคากลางเพื่อประเมินค่าซ่อมเบื้องต้นโดยอัตโนมัติ", prefix="Auto Cost Estimation — ")

    doc.add_page_break()

    # ------------------------------------------------------------------
    # PAGE 6 — 1.5 + 1.6
    # ------------------------------------------------------------------
    sub_heading("1.5   ระบบตรวจสภาพรถด้วยตนเองผ่านลิงก์")
    body("สำหรับกรณีต่ออายุกรมธรรม์ที่ต้องตรวจสภาพรถก่อนคุ้มครอง:", after=4)
    bullet("เจ้าหน้าที่สร้างเคสตรวจสภาพ ระบบส่งอีเมลพร้อมลิงก์เฉพาะเคสให้ลูกค้าถ่ายภาพเองผ่านมือถือ ไม่ต้องล็อกอิน", prefix="Unique Token Link — ")
    bullet("แนะนำให้ลูกค้าถ่ายภาพรถ 8 มุมมาตรฐาน พร้อมภาพเล่มทะเบียน/เลขตัวถัง", prefix="8-Angle Photo Capture — ")
    bullet("AI ตรวจสอบความครบถ้วนและคุณภาพของภาพก่อนส่งให้เจ้าหน้าที่ตรวจสอบต่อ ลดปัญหาภาพเบลอ/มุมไม่ครบ", prefix="AI Photo Quality Check — ")

    sub_heading("1.6   แคตตาล็อกราคากลางและการนำเข้าข้อมูล")
    body("ฐานข้อมูลราคาอ้างอิงสำหรับควบคุมราคาซ่อม:", after=4)
    bullet("ราคากลางค่าแรงซ่อม/พ่นสี/เคาะ ตามมาตรฐานอ้างอิง แยกตามยี่ห้อ ขนาดรถ และระดับความเสียหาย", prefix="Standard Repair Labor Catalog — ")
    bullet("ฐานราคาอะไหล่แยกตามยี่ห้อ รุ่น ปี และประเภทอะไหล่ (แท้ / เทียบเท่า / มือสอง)", prefix="Spare Parts Price Catalog — ")
    bullet("นำเข้า/อัปเดตราคาจำนวนมากผ่านไฟล์ Excel พร้อมจับคู่หัวคอลัมน์ไทย/อังกฤษอัตโนมัติ และบันทึกประวัติการนำเข้า", prefix="Bulk Excel Import — ")

    doc.add_page_break()

    # ------------------------------------------------------------------
    # PAGE 7 — 1.7 + 1.8 + 1.9
    # ------------------------------------------------------------------
    sub_heading("1.7   ระบบบริหารจัดการส่วนกลาง")
    body("ศูนย์ควบคุมสำหรับผู้ดูแลระบบ:", after=4)
    bullet("เพิ่ม/แก้ไข/ปิดใช้งานสาขาได้เองโดยไม่ต้องแก้โค้ด รองรับการขยายสาขาใหม่ในอนาคต", prefix="Branch Management — ")
    bullet("กำหนดบทบาทและสิทธิ์การเข้าถึงแต่ละฟังก์ชันได้อย่างละเอียด", prefix="Role Management — ")
    bullet("เพิ่มพนักงานทีละคนหรือนำเข้าจำนวนมากผ่านไฟล์ Excel", prefix="User Management — ")
    bullet("ตั้งค่าเงื่อนไขการอนุมัติ เช่น วงเงินที่ต้องผ่านผู้มีอำนาจอนุมัติ และรายชื่อบทบาทที่มีสิทธิ์อนุมัติแต่ละระดับ", prefix="Workflow Configuration — ")

    sub_heading("1.8   รายงานและสถิติเชิงบริหาร")
    body("มุมมองสำหรับผู้บริหารในการติดตามผลการดำเนินงาน:", after=4)
    bullet("ติดตามระยะเวลาการอนุมัติเคลมเทียบกับเป้าหมาย SLA ที่ตั้งไว้", prefix="SLA Turnaround Report — ")
    bullet("จัดอันดับอู่/ศูนย์บริการตามอัตราการเสนอราคาสูงกว่าราคากลาง ช่วยระบุคู่ค้าที่ควรเฝ้าระวัง", prefix="Garage Integrity Score — ")
    bullet("สรุปยอดเคลมรวม ยอดอนุมัติสุทธิ และยอดประหยัดสะสมจากการควบคุมราคา", prefix="Executive Dashboard — ")

    sub_heading("1.9   ระบบบันทึกประวัติการทำงาน")
    body("ทุกการเปลี่ยนแปลงสำคัญถูกบันทึกไว้เพื่อการตรวจสอบย้อนหลัง ทั้งการเปลี่ยนสถานะเคส (Quotation Log) "
         "และการเปลี่ยนแปลงข้อมูลระดับผู้ดูแลระบบ เช่น สาขา บทบาท ผู้ใช้งาน และเงื่อนไขการอนุมัติ (System Audit Log) "
         "พร้อมระบุผู้ทำรายการและเวลาที่เกิดขึ้นครบถ้วน", after=4)

    doc.add_page_break()

    # ------------------------------------------------------------------
    # PAGE 8 — CHAPTER 2 ARCHITECTURE
    # ------------------------------------------------------------------
    chapter_banner("02", "สถาปัตยกรรมระบบและ API ที่ใช้งาน", "System Architecture & APIs")
    body("ระบบพัฒนาบนสถาปัตยกรรม Full-Stack Web Application ทำงานบน VPS ของบริษัทเอง "
         "ควบคุมได้เต็มรูปแบบทั้งด้านข้อมูลและความปลอดภัย โดยเชื่อมต่อกับบริการภายนอกเท่าที่จำเป็นดังนี้:", after=8)

    data_table(
        ["Layer / โมดูล", "API / เทคโนโลยีที่ใช้", "หน้าที่หลัก"],
        [
            ("AI Document & Vision Engine", "Anthropic Claude API\n(Vision + Text)",
             "ถอดข้อมูลจากใบเสนอราคา (PDF/รูปภาพ) วิเคราะห์ภาพความเสียหายรถยนต์ ตรวจสอบคุณภาพภาพตรวจสภาพรถ"),
            ("Native PDF Text Parser", "pdf-parse (Node.js)",
             "ถอดข้อความจาก PDF ที่มีเลเยอร์ข้อความโดยตรง ก่อนเรียกใช้ AI เพื่อลดต้นทุนและเพิ่มความเร็ว"),
            ("Image Processing", "Sharp",
             "ย่อ/บีบอัดรูปภาพก่อนส่งวิเคราะห์ด้วย AI Vision เพื่อลดขนาดข้อมูลและควบคุมต้นทุน"),
            ("Transactional Email", "Resend API",
             "ส่งอีเมลลิงก์ตรวจสภาพรถให้ลูกค้า และการแจ้งเตือนสถานะต่างๆ เป็นภาษาไทย"),
            ("Application Database", "Prisma ORM + SQLite\n(better-sqlite3)",
             "จัดเก็บข้อมูลเคส แคตตาล็อกราคา ผู้ใช้งาน สาขา และประวัติการทำงาน ด้วย Prepared Statements ป้องกัน SQL Injection"),
        ],
        widths=[Cm(4.4), Cm(4.0), Cm(8.0)],
    )

    callout([
        "• Direct Download Block — บล็อกการเข้าถึงไฟล์ .xlsx, .xls, .csv, .db ผ่าน URL ตรงทั้งหมด",
        "• Data Directory Isolation — ไฟล์ข้อมูลธุรกิจและฐานข้อมูลจัดเก็บนอก Web Root ไม่สามารถเข้าถึงผ่านเว็บได้",
        "• Zero-Downtime Deployment — ปรับปรุงระบบ Production ผ่าน PM2 Process Manager โดยไม่หยุดให้บริการ",
    ], title="สถาปัตยกรรมความปลอดภัยของข้อมูล")

    doc.add_page_break()

    # ------------------------------------------------------------------
    # PAGE 9 — CHAPTER 3 STACK
    # ------------------------------------------------------------------
    chapter_banner("03", "เทคโนโลยีทั้งหมดที่ใช้ในโปรเจกต์", "Technology Stack")
    body("รายละเอียดเทคโนโลยีหลักที่ใช้พัฒนาระบบ พร้อมเหตุผลในการเลือกใช้แต่ละส่วน:", after=8)

    data_table(
        ["หมวดหมู่", "เทคโนโลยี / ไลบรารี", "เหตุผลในการเลือกใช้"],
        [
            ("Web Application Framework", "Next.js (App Router)\nReact + TypeScript",
             "ประสิทธิภาพสูง โหลดหน้าเร็ว รองรับ SSR/CSR พร้อมระบบตรวจชนิดข้อมูลลดข้อผิดพลาด"),
            ("Styling & UI", "Tailwind CSS\nCustom Design System",
             "ดีไซน์พรีเมียมโทนสีน้ำเงิน (#0071e3) ตอบสนองทุกขนาดหน้าจอ"),
            ("Database & ORM", "Prisma ORM\nSQLite (better-sqlite3)",
             "จัดการฐานข้อมูลปลอดภัยด้วย Prepared Statements ติดตั้งง่าย ไม่ต้องพึ่งเซิร์ฟเวอร์ฐานข้อมูลแยก"),
            ("AI Engine", "Anthropic Claude API\n(@anthropic-ai/sdk)",
             "โมเดล AI อ่านเอกสารและวิเคราะห์ภาพ ความแม่นยำสูงสำหรับภาษาไทย"),
            ("Document / Spreadsheet", "xlsx, pdf-parse, adm-zip",
             "นำเข้า/ส่งออกข้อมูลแคตตาล็อกราคาผ่าน Excel และถอดข้อความจากไฟล์ PDF"),
            ("Image Processing", "Sharp",
             "ปรับขนาดและบีบอัดรูปภาพก่อนประมวลผลด้วย AI เพื่อควบคุมต้นทุนและความเร็ว"),
            ("Email Services", "Resend API",
             "ส่งอีเมลแจ้งเตือนสถานะและลิงก์ตรวจสภาพรถถึงลูกค้าเป็นภาษาไทยโดยอัตโนมัติ"),
            ("Deployment", "PM2 + Node.js LTS\nบน VPS ของบริษัท",
             "ควบคุมการรันแอปพลิเคชันบน Production พร้อม Auto-Restart และ Zero-Downtime Deployment"),
        ],
        widths=[Cm(4.4), Cm(4.0), Cm(8.0)],
    )

    doc.add_page_break()

    # ------------------------------------------------------------------
    # PAGE 10 — CHAPTER 4 PRICING (intro + tiers)
    # ------------------------------------------------------------------
    chapter_banner("04", "ข้อเสนอราคา", "Pay-Per-Completed-Transaction Model")
    body("บริษัท เทคทันใจ อินโนเวชั่น จำกัด นำเสนอโมเดลค่าบริการที่คิดเฉพาะ "
         "“เคสที่ผ่านกระบวนการสแกน ควบคุมราคา และอนุมัติเสร็จสมบูรณ์ (Completed Transaction)” เท่านั้น:", after=6)
    bullet("เคสที่เป็นแบบร่าง ถูกยกเลิก หรือถูกปฏิเสธการอนุมัติ จะไม่มีการคิดค่าบริการใดๆ ทั้งสิ้น", prefix="ความเสี่ยงทางธุรกิจต่ำที่สุด — ")
    bullet("ราคารวมค่าใช้จ่ายเซิร์ฟเวอร์ ค่า AI API และการดูแลรักษาระบบไว้ในราคาเดียวแล้ว ไม่มีค่าใช้จ่ายแอบแฝง", prefix="ราคาแบบเหมารวม — ")
    bullet("ยิ่งใช้งานปริมาณมากขึ้น ราคาต่อเคสยิ่งลดลง เหมาะกับการเติบโตของธุรกิจในระยะยาว", prefix="ส่วนลดตามปริมาณการใช้งาน — ")

    sub_heading("4.1   ตารางอัตราค่าบริการตามปริมาณการใช้งาน")
    data_table(
        ["แพ็กเกจ", "ปริมาณเคส/เดือน", "ราคาต่อเคสสำเร็จ", "เหมาะสำหรับ"],
        [
            ("Starter", "1 – 500 เคส", "30 บาท", "อู่ซ่อม / สาขาเดี่ยว เริ่มทดลองใช้งาน"),
            ("Growth", "501 – 1,000 เคส", "27 บาท", "เครือข่ายอู่ซ่อมหรือหน่วยงานหลายสาขา"),
            ("Enterprise", "1,001 – 3,000 เคส", "25 บาท", "บริษัทประกันภัย / TPA ขนาดกลาง-ใหญ่"),
            ("Strategic Partner", "3,000+ เคส", "22 บาท", "องค์กรประกันภัยระดับประเทศ"),
        ],
        widths=[Cm(3.6), Cm(3.6), Cm(4.2), Cm(5.0)],
        header_color="0A2540",
    )

    doc.add_page_break()

    # ------------------------------------------------------------------
    # PAGE 11 — 4.2 EXAMPLE + CONTACT
    # ------------------------------------------------------------------
    sub_heading("4.2   ตัวอย่างประมาณการที่ปริมาณ 1,000 เคส/เดือน")
    data_table(
        ["รายการ", "มูลค่าโดยประมาณ"],
        [
            ("ยอดเสนอราคาเคลมรวม (1,000 เคส x 15,000 บาท/เคส โดยประมาณ)", "15,000,000 บาท / เดือน"),
            ("ค่าบริการระบบ Claim-Thunjai (Growth: 1,000 เคส x 27 บาท)", "27,000 บาท / เดือน"),
            ("สัดส่วนค่าบริการต่อยอดเคลมรวม", "ต่ำกว่า 0.18% ของยอดเคลม"),
        ],
        widths=[Cm(10.8), Cm(5.6)],
        header_color="0F4C81",
    )
    doc.add_paragraph().paragraph_format.space_after = Pt(2)
    body("หมายเหตุ: ตัวเลขยอดเคลมรวมและอัตราการประหยัดที่แท้จริงขึ้นอยู่กับข้อมูลปริมาณเคสและพฤติกรรมการเสนอราคาของอู่ซ่อมของแต่ละองค์กร "
         "ทางบริษัทยินดีจัดทำประมาณการเฉพาะสำหรับองค์กรของท่านจากข้อมูลจริง", italic=True, after=10)

    callout([
        "บริษัท เทคทันใจ อินโนเวชั่น จำกัด พร้อมสาธิตระบบใช้งานจริง (Live Demonstration) และปรับแผนราคาให้เหมาะสมกับปริมาณการใช้งานขององค์กร",
        "",
        "• โทรศัพท์: 065-882-8333    Email: athaporn@techthunjai.com    Line Official: @techthunjai",
        "• ระบบทดลองใช้งานจริง: https://demo-claim.techthunjai.com/",
    ], title="ติดต่อทีมงาน")

    # ------------------------------------------------------------------
    # FOOTER — page numbers on every page
    # ------------------------------------------------------------------
    def add_page_number_footer(section):
        footer = section.footer
        footer.is_linked_to_previous = False
        p = footer.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        pPr = p._p.get_or_add_pPr()
        pPr.append(parse_xml(
            r'<w:pBdr xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">'
            r'<w:top w:val="single" w:sz="6" w:space="6" w:color="D7E2EF"/></w:pBdr>'))

        r0 = p.add_run("บริษัท เทคทันใจ อินโนเวชั่น จำกัด    ")
        set_font(r0, size=10, color=MUTED)

        r1 = p.add_run()
        fld_begin = OxmlElement('w:fldChar')
        fld_begin.set(qn('w:fldCharType'), 'begin')
        r1._r.append(fld_begin)

        r2 = p.add_run()
        instr = OxmlElement('w:instrText')
        instr.set(qn('xml:space'), 'preserve')
        instr.text = ' PAGE '
        r2._r.append(instr)

        r3 = p.add_run()
        fld_sep = OxmlElement('w:fldChar')
        fld_sep.set(qn('w:fldCharType'), 'separate')
        r3._r.append(fld_sep)

        r4 = p.add_run("1")
        set_font(r4, size=11, bold=True, color=BLUE)

        r5 = p.add_run()
        fld_end = OxmlElement('w:fldChar')
        fld_end.set(qn('w:fldCharType'), 'end')
        r5._r.append(fld_end)

        for run in (r0, r4):
            pass  # already styled above

    for section in doc.sections:
        add_page_number_footer(section)

    os.makedirs("data", exist_ok=True)
    file_path = os.path.abspath("data/Claude_Proposal_Claim_Thunjai.docx")
    doc.save(file_path)
    print("SUCCESS:", file_path)
    return file_path


if __name__ == "__main__":
    create_proposal_docx()
