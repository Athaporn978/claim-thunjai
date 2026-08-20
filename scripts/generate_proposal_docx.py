import os
import docx
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import nsdecls, qn

def create_proposal_docx():
    doc = Document()

    # Page setup - Standard A4 with 1 inch margins
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

    # Base Styles
    COLOR_PRIMARY = RGBColor(0, 113, 227)     # #0071e3 Premium Vibrant Blue
    COLOR_DARK_BLUE = RGBColor(15, 76, 129)   # #0f4c81 Deep Navy
    COLOR_TEXT = RGBColor(51, 51, 51)         # #333333 Dark Slate Text
    COLOR_MUTED = RGBColor(100, 116, 139)     # #64748b Slate
    COLOR_WHITE = RGBColor(255, 255, 255)

    FONT_FAMILY = "TH Sarabun New"

    # Helper function to style paragraph
    def format_p(p, font_size=16, bold=False, italic=False, color=COLOR_TEXT, align=WD_ALIGN_PARAGRAPH.LEFT, space_before=0, space_after=6, line_spacing=1.15):
        p.alignment = align
        p.paragraph_format.space_before = Pt(space_before)
        p.paragraph_format.space_after = Pt(space_after)
        p.paragraph_format.line_spacing = line_spacing
        for run in p.runs:
            run.font.name = FONT_FAMILY
            run.font.size = Pt(font_size)
            run.bold = bold
            run.italic = italic
            run.font.color.rgb = color
            # Set East Asia font for Thai compatibility
            rPr = run._r.get_or_add_rPr()
            rFonts = OxmlElement('w:rFonts')
            rFonts.set(qn('w:ascii'), FONT_FAMILY)
            rFonts.set(qn('w:hAnsi'), FONT_FAMILY)
            rFonts.set(qn('w:cs'), FONT_FAMILY)
            rFonts.set(qn('w:eastAsia'), FONT_FAMILY)
            rPr.append(rFonts)

    def add_styled_heading(text, level=1):
        p = doc.add_paragraph()
        run = p.add_run(text)
        if level == 1:
            format_p(p, font_size=24, bold=True, color=COLOR_PRIMARY, space_before=18, space_after=8)
            pPr = p._p.get_or_add_pPr()
            pBdr = parse_xml(r'<w:pBdr xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">'
                             r'<w:bottom w:val="single" w:sz="12" w:space="4" w:color="0071E3"/>'
                             r'</w:pBdr>')
            pPr.append(pBdr)
        elif level == 2:
            format_p(p, font_size=20, bold=True, color=COLOR_DARK_BLUE, space_before=14, space_after=6)
        elif level == 3:
            format_p(p, font_size=18, bold=True, color=COLOR_PRIMARY, space_before=10, space_after=4)
        return p

    def add_body_p(text, bold=False, italic=False, space_after=6):
        p = doc.add_paragraph()
        p.add_run(text)
        format_p(p, font_size=16, bold=bold, italic=italic, color=COLOR_TEXT, space_after=space_after)
        return p

    def add_bullet_p(text, bold_prefix="", space_after=4):
        p = doc.add_paragraph(style='List Bullet')
        if bold_prefix:
            run_b = p.add_run(bold_prefix)
            run_b.bold = True
            run_b.font.color.rgb = COLOR_DARK_BLUE
        run_t = p.add_run(text)
        format_p(p, font_size=16, color=COLOR_TEXT, space_after=space_after)
        return p

    def set_cell_background(cell, fill_hex):
        shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{fill_hex}"/>')
        cell._tc.get_or_add_tcPr().append(shading)

    def add_callout_box(text_list, title=""):
        tbl = doc.add_table(rows=1, cols=1)
        tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
        cell = tbl.cell(0, 0)
        set_cell_background(cell, "F0F7FF") # Light Soft Blue tint
        
        tcPr = cell._tc.get_or_add_tcPr()
        tcBorders = parse_xml(r'<w:tcBorders xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">'
                              r'<w:top w:val="none"/>'
                              r'<w:left w:val="single" w:sz="36" w:color="0071E3"/>'
                              r'<w:bottom w:val="none"/>'
                              r'<w:right w:val="none"/>'
                              r'</w:tcBorders>')
        tcPr.append(tcBorders)

        p = cell.paragraphs[0]
        if title:
            p.add_run(f"💡 {title}\n").bold = True
        for idx, line in enumerate(text_list):
            if idx > 0 or title:
                p = cell.add_paragraph()
            p.add_run(line)
            format_p(p, font_size=15, color=COLOR_DARK_BLUE, space_after=4)
        doc.add_paragraph() # Spacer

    # ==========================================
    # 1. หน้าปก (COVER PAGE)
    # ==========================================
    p_top_spacer = doc.add_paragraph()
    format_p(p_top_spacer, space_before=40)

    p_org = doc.add_paragraph()
    p_org.add_run("บริษัท เทคทันใจ อินโนเวชั่น จำกัด\nTECHTHUNJAI INNOVATION CO., LTD.")
    format_p(p_org, font_size=18, bold=True, color=COLOR_MUTED, align=WD_ALIGN_PARAGRAPH.RIGHT, space_after=40)

    p_badge = doc.add_paragraph()
    p_badge.add_run("EXECUTIVE PROJECT PROPOSAL")
    format_p(p_badge, font_size=14, bold=True, color=COLOR_PRIMARY, space_after=10)

    p_main_title = doc.add_paragraph()
    p_main_title.add_run("ข้อเสนอโครงการระบบบริหารจัดการเคลม\nและคุมราคาไมโครเอไออัจฉริยะ")
    format_p(p_main_title, font_size=28, bold=True, color=COLOR_PRIMARY, space_after=12)

    p_sub_title = doc.add_paragraph()
    p_sub_title.add_run("Claim-Thunjai AI Platform: Intelligent Auto-Quote Extraction & Claim Analytics System")
    format_p(p_sub_title, font_size=18, italic=True, color=COLOR_DARK_BLUE, space_after=60)

    # Decorative colored block line XML
    p_line = doc.add_paragraph()
    pPr = p_line._p.get_or_add_pPr()
    pBdr = parse_xml(r'<w:pBdr xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">'
                     r'<w:bottom w:val="single" w:sz="36" w:space="1" w:color="0071E3"/>'
                     r'</w:pBdr>')
    pPr.append(pBdr)

    p_info = doc.add_paragraph()
    p_info.add_run("นำเสนอสำหรับ: ").bold = True
    p_info.add_run("คณะผู้บริหารระดับสูง (CEO & Board of Directors)\n")
    p_info.add_run("จัดทำโดย: ").bold = True
    p_info.add_run("บริษัท เทคทันใจ อินโนเวชั่น จำกัด (TECHTHUNJAI INNOVATION CO., LTD.)\n")
    p_info.add_run("ช่องทางติดต่อ: ").bold = True
    p_info.add_run("Email: athaporn@techthunjai.com | Line OA: @techthunjai | Tel: 065-882-8333\n")
    p_info.add_run("วันที่นำเสนอ: ").bold = True
    p_info.add_run("ปีพุทธศักราช 2569 (2026)")
    format_p(p_info, font_size=16, color=COLOR_TEXT, space_before=20, space_after=20)

    doc.add_page_break()

    # ==========================================
    # 2. สารบัญ (TABLE OF CONTENTS)
    # ==========================================
    add_styled_heading("สารบัญ (Table of Contents)", level=1)
    
    toc_items = [
        ("บทสรุปผู้บริหาร (Executive Summary)", "หน้า 3"),
        ("หมวดที่ 1: รายละเอียดโมดูลและฟีเจอร์ทั้งหมด (System Modules & Features)", "หน้า 4"),
        ("   1.1 โมดูลอ่านเอกสารใบเสนอราคาอัจฉริยะ (AI Auto-Quote Extraction)", "หน้า 4"),
        ("   1.2 โมดูลตรวจวิเคราะห์รอยแผลความเสียหายจากภาพถ่าย (AI Visual Damage Inspection)", "หน้า 4"),
        ("   1.3 โมดูลคุมราคาเปรียบเทียบราคากลางอู่และประกันภัย (Dynamic Price Control Catalog)", "หน้า 5"),
        ("   1.4 โมดูลบริหารจัดการสาขาและลำดับการอนุมัติ (Multi-Branch RBAC & Workflows)", "หน้า 5"),
        ("   1.5 โมดูลรายงานวิเคราะห์เชิงบริหารและสถิติ (Executive Dashboard & Analytics)", "หน้า 6"),
        ("   1.6 โมดูลดักรับอีเมลและเชื่อมต่อระบบภายนอก (Gmail Poller & Intake Integration)", "หน้า 6"),
        ("หมวดที่ 2: สถาปัตยกรรมระบบและ API ทั้งหมด (System Architecture & API Engines)", "หน้า 7"),
        ("หมวดที่ 3: เทคโนโลยีทั้งหมดที่ใช้ในโปรเจกต์ (Complete Technology Stack)", "หน้า 8"),
        ("หมวดที่ 4: ข้อเสนอราคาและโมเดลการคิดค่าบริการ (Pricing Proposal & Business Model)", "หน้า 9"),
        ("   4.1 โมเดลการคิดค่าบริการแบบ Pay-Per-Completed-Transaction", "หน้า 9"),
        ("   4.2 ตารางอัตราค่าบริการตามปริมาณการใช้งาน (Volume Tiering Table)", "หน้า 9"),
        ("   4.3 ประมาณการผลตอบแทนและการประหยัดต้นทุน (ROI Savings Analysis)", "หน้า 10"),
    ]

    for title, page in toc_items:
        p = doc.add_paragraph()
        p.add_run(title)
        format_p(p, font_size=16, color=COLOR_DARK_BLUE if "หมวด" in title else COLOR_TEXT, space_after=4)

    doc.add_page_break()

    # Executive Summary
    add_styled_heading("บทสรุปผู้บริหาร (Executive Summary)", level=1)
    add_body_p("ระบบ Claim-Thunjai AI Platform ถูกออกแบบและพัฒนาโดย บริษัท เทคทันใจ อินโนเวชั่น จำกัด เพื่อปฏิวัติกระบวนการอนุมัติเคลมประกันภัยรถยนต์และการคุมราคาซ่อมอู่กลาง/ศูนย์บริการ โดยเปลี่ยนกระบวนการเดิมที่ต้องใช้เจ้าหน้าที่คุมราคาตรวจทานเอกสารทีละหน้า ให้กลายเป็นกระบวนการอัตโนมัติด้วย AI 100%")
    
    add_callout_box([
        "• ลดระยะเวลาการอนุมัติเคลมจากเฉลี่ย 3-5 วัน เหลือเพียง 30 วินาทีต่อเคส",
        "• ประหยัดงบประมาณการจ่ายเคลมส่วนเกิน (Overcharge Claim) ได้เฉลี่ย 12-18% ของยอดเคลมรวม",
        "• รองรับการทำงานแบบ Dynamic Branch RBAC แยกการมองเห็นตามสาขาและสิทธิ์การอนุมัติ 3 ระดับ",
        "• คิดค่าบริการเฉพาะเคสที่อนุมัติสำเร็จ 100% (Completed Transaction Model) เสี่ยงเป็นศูนย์สำหรับธุรกิจ"
    ], title="จุดเด่นและผลประโยชน์หลักทางธุรกิจ (Key Business Value)")

    # ==========================================
    # 3. MODULES & FEATURES
    # ==========================================
    add_styled_heading("หมวดที่ 1: รายละเอียดโมดูลและฟีเจอร์ทั้งหมด (Modules & Features)", level=1)

    add_styled_heading("1.1 โมดูลอ่านเอกสารใบเสนอราคาอัจฉริยะ (AI Auto-Quote Extraction)", level=2)
    add_body_p("โมดูล AI สำหรับสแกนและถอดข้อมูลจากเอกสารใบเสนอราคาซ่อม (Repair Quotation) ทั้งไฟล์ PDF และรูปถ่ายภาพสแกน เข้าสู่ฟอร์มระบบอัตโนมัติ 6 ขั้นตอน:")
    add_bullet_p("สกัดชื่อ-นามสกุล, เบอร์โทรศัพท์, ที่อยู่ และเลขประจำตัวผู้เสียภาษีอัตโนมัติ", bold_prefix="1. ข้อมูลลูกค้า (Customer Metadata): ")
    add_bullet_p("สกัดชื่อบริษัทประกันภัย, เลขกรมธรรม์, เลขที่ใบรับแจ้ง/เลขเคลม และประเภทความคุ้มครอง", bold_prefix="2. ข้อมูลประกันภัย (Insurance Detail): ")
    add_bullet_p("สกัดชื่อศูนย์บริการ/อู่ซ่อม, เบอร์โทรศัพท์ติดต่อ และที่อยู่อู่สังกัด", bold_prefix="3. ศูนย์บริการ/อู่ซ่อม (Repair Center): ")
    add_bullet_p("จัดเก็บและแสดงผลไฟล์ภาพถ่ายใบเสนอราคาต้นฉบับเพื่อการตรวจสอบยันกลับ", bold_prefix="4. ภาพถ่ายเอกสาร (Document Photos): ")
    add_bullet_p("แยกรายการซ่อมประเภทค่าแรง พ่นสี ถอดประกอบ เคาะซ่อม และคำนวณเปรียบเทียบราคากลาง", bold_prefix="5. รายการค่าแรง (Labor Items): ")
    add_bullet_p("แยกรายการชิ้นส่วนอะไหล่ รหัสอะไหล่ จำนวน และราคาต่อหน่วย พร้อมคำนวณส่วนลดสุทธิ", bold_prefix="6. รายการค่าอะไหล่ (Part Items): ")

    add_styled_heading("1.2 โมดูลตรวจวิเคราะห์รอยแผลความเสียหายจากภาพถ่าย (AI Visual Damage Inspection)", level=2)
    add_body_p("โมดูลปัญญาประดิษฐ์ประมวลผลภาพถ่าย (Computer Vision) ตรวจจับรอยบาดแผลและประเมินระดับความเสียหายของชิ้นส่วนรถยนต์:")
    add_bullet_p("ระบุมุมมองภาพถ่าย (Front, Rear, Left Door, Fender, Roof) อัตโนมัติ", bold_prefix="• Multi-Angle Photo Detection: ")
    add_bullet_p("ตรวจจับตำแหน่งรอยขีดข่วน (Scratch), รอยบุบ (Dent), รอยถลอก และรอยชนหนัก พร้อมระบุ Bounding Box (%)", bold_prefix="• Surgical Bounding Box Detection: ")
    add_bullet_p("จัดระดับความเสียหาย 4 ระดับ (Minor, Moderate, Severe, Total Loss) พร้อมประมาณการราคาค่าซ่อมเบื้องต้น", bold_prefix="• Severity Ranking & Price Estimation: ")

    add_styled_heading("1.3 โมดูลคุมราคาเปรียบเทียบราคากลาง (Dynamic Price Control Catalog)", level=2)
    add_body_p("ระบบฐานข้อมูลราคากลางและการคุมราคาอัจฉริยะ (Fair Matching Control Engine):")
    add_bullet_p("บรรจุรายการราคากลางค่าแรงซ่อม พ่นสี เคาะ ถอดประกอบ มาตรฐาน 504 รายการหลัก ครอบคลุมรถยนต์ทุกยี่ห้อ (Toyota, Honda, Nissan, Isuzu, Mazda, MG, BYD, BMW, Benz)", bold_prefix="• Master Repair Price Catalog: ")
    add_bullet_p("หากราคาที่อู่เสนอมา สูงกว่าราคากลาง ระบบจะปรับลดราคาควบคุม (Controlled Price) ลงมาเท่าราคากลางทันที แต่หากราคาอู่น้อยกว่าราคากลาง ระบบจะอนุมัติราคาตามจริง 100% (ไม่มีการปรับราคาขึ้น)", bold_prefix="• Fair Matching Rule Engine: ")
    add_bullet_p("ช่องส่วนลดเปอร์เซ็นต์ (% Discount) และจำนวนเงินส่วนลด สามารถพิมพ์แก้ไขได้อย่างอิสระโดยไม่โดนตัวเลขขยะทับ", bold_prefix="• Flexible Discount Input: ")

    add_styled_heading("1.4 โมดูลบริหารจัดการสาขาและลำดับการอนุมัติ (Multi-Branch RBAC & Workflows)", level=2)
    add_body_p("ระบบจัดการสิทธิ์ผู้ใช้งานและการอนุมัติหลายระดับตามโครงสร้างองค์กร:")
    add_bullet_p("ผูกผู้สร้างเคสและข้อมูลเข้ากับสาขาสังกัดแบบ Dynamic 100% (รองรับสาขาใหม่ ชลบุรี, ขอนแก่น ฯลฯ โดยไม่ต้องแก้โค้ด)", bold_prefix="• Dynamic Branch Binding: ")
    add_bullet_p("รองรับบทบาท Super Admin, หัวหน้าคุมราคา (Supervisor), เจ้าหน้าที่คุมราคา (Officer) และอู่ซ่อม/ศูนย์บริการ", bold_prefix="• Multi-Role Permission (RBAC): ")
    add_bullet_p("อนุมัติเคสตามวงเงินสิทธิ์ หากเกินวงเงินระบบจะส่งต่อไปยังหัวหน้าสาขาหรือผู้จัดการโดยอัตโนมัติ", bold_prefix="• 3-Tier Approval Flow: ")

    add_styled_heading("1.5 โมดูลรายงานวิเคราะห์เชิงบริหารและสถิติ (Executive Dashboard & Analytics)", level=2)
    add_body_p("ระบบรายงานสรุปผลการดำเนินงานแบบเรียลไทม์สำหรับระดับผู้บริหาร:")
    add_bullet_p("สรุปยอดเคลมรวม, ยอดอนุมัติสุทธิ, ยอดประหยัดสะสม (Net Savings) และเปอร์เซ็นต์การประหยัด", bold_prefix="• Executive Summary Cards: ")
    add_bullet_p("วิเคราะห์สัดส่วนค่าแรง vs ค่าอะไหล่ สรุปผลรวมเป็น 100% เพื่อความชัดเจนในการบริหารจัดการ", bold_prefix="• Proportional Claim Breakdown Chart: ")
    add_bullet_p("ติดตามระยะเวลาการอนุมัติเคลมเทียบกับเป้าหมาย SLA (Service Level Agreement)", bold_prefix="• SLA & Productivity Tracking: ")

    add_styled_heading("1.6 โมดูลดักรับอีเมลและเชื่อมต่อระบบภายนอก (Gmail Poller & Intake Integration)", level=2)
    add_body_p("ระบบเชื่อมต่อการรับเอกสารเคลมจากช่องทางภายนอกอัตโนมัติ:")
    add_bullet_p("คอยตรวจจับอีเมลที่มีไฟล์แนบใบเสนอราคา (.pdf, .jpg, .png) จากอู่ซ่อม และนำเข้าสู่ระบบสแกนอัตโนมัติ", bold_prefix="• IMAP Automated Email Poller: ")
    add_bullet_p("รับข้อมูลเคสเคลมจากระบบ Core Insurance หรือระบบเคลมภายนอกผ่าน RESTful API", bold_prefix="• Third-Party REST API Intake: ")

    doc.add_page_break()

    # ==========================================
    # 4. SYSTEM ARCHITECTURE & APIS
    # ==========================================
    add_styled_heading("หมวดที่ 2: สถาปัตยกรรมระบบและ API ทั้งหมด (System Architecture)", level=1)
    add_body_p("ระบบ Claim-Thunjai AI Platform ถูกออกแบบตามสถาปัตยกรรม Modern Micro-Service & Serverless Hybrid Architecture เพื่อความเสถียรสูงสุด ความเร็วในการประมวลผลระดับมิลลิวินาที และความปลอดภัยข้อมูลระดับสูงสุด:")

    # Architecture Table
    arch_tbl = doc.add_table(rows=5, cols=3)
    arch_tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    headers = ["Layer / Module", "API / Technology Used", "หน้าที่และความสามารถหลัก (Core Responsibility)"]
    hdr_cells = arch_tbl.rows[0].cells
    for i, title in enumerate(headers):
        hdr_cells[i].text = title
        set_cell_background(hdr_cells[i], "0071E3")
        format_p(hdr_cells[i].paragraphs[0], font_size=15, bold=True, color=COLOR_WHITE, align=WD_ALIGN_PARAGRAPH.CENTER)

    arch_data = [
        ("AI Document Vision Engine", "Anthropic Claude 3.5 Sonnet Vision API\n(PDFs & Multimodal Vision)", "วิเคราะห์และถอดโครงสร้างเอกสารใบเสนอราคา ภาษาไทย/อังกฤษ และรูปภาพสแกนอย่างแม่นยำ"),
        ("High-Speed AI Processing Engine", "Google Gemini 2.5 / Flash AI Engine", "ประมวลผลข้อมูลเอกสารขนาดใหญ่ความเร็วสูง และวิเคราะห์รอยแผลบาดแผลรถยนต์ (Image Detection)"),
        ("Native Text & Vector Parser", "Node.js Native pdf-parse & Sharp Engine", "ถอดข้อความ vector จากไฟล์ PDF ดิจิทัลโดยตรงความเร็วสูง 100% โดยไม่ต้องพึ่งพา AI ภายนอก"),
        ("Email Poller & Messaging API", "Gmail IMAP Flow Engine & Resend Email API", "ดักรับอีเมลใบเสนอราคาจากอู่ซ่อมอัตโนมัติ และส่งอีเมลการอนุมัติเคลม/แจ้งเตือนสถานะอนุมัติ"),
    ]

    for row_idx, data in enumerate(arch_data, start=1):
        row_cells = arch_tbl.rows[row_idx].cells
        if row_idx % 2 == 1:
            for c in row_cells: set_cell_background(c, "F9FAFB")
        row_cells[0].text = data[0]
        row_cells[1].text = data[1]
        row_cells[2].text = data[2]
        format_p(row_cells[0].paragraphs[0], font_size=14, bold=True, color=COLOR_DARK_BLUE)
        format_p(row_cells[1].paragraphs[0], font_size=14, color=COLOR_PRIMARY)
        format_p(row_cells[2].paragraphs[0], font_size=14, color=COLOR_TEXT)

    doc.add_paragraph() # Spacer

    add_callout_box([
        "• Security Standard: ระบบรองรับการบล็อกการดาวน์โหลดไฟล์ตรง (Direct Download Block) สำหรับไฟล์ .xlsx, .db, .csv จากภายนอก 100%",
        "• Data Directory Isolation: ไฟล์ข้อมูลธุรกิจและแคตตาล็อกราคากลางทั้งหมดจัดเก็บในโฟลเดอร์ปิด data/ นอก Web Root ป้องกันการรั่วไหล 100%",
        "• Zero-Downtime Deployment: ระบบรองรับการอัปเดตโค้ดแบบไร้การหยุดชะงักผ่าน PM2 Cluster Mode บน VPS Production"
    ], title="สถาปัตยกรรมความปลอดภัยและการจัดเก็บข้อมูล (Data Security Architecture)")

    doc.add_page_break()

    # ==========================================
    # 5. TECH STACK
    # ==========================================
    add_styled_heading("หมวดที่ 3: เทคโนโลยีทั้งหมดที่ใช้ในโปรเจกต์ (Technology Stack)", level=1)
    add_body_p("รายละเอียดเทคโนโลยีระดับ Enterprise Grade ที่ใช้ในการพัฒนาระบบทั้งหมด:")

    stack_tbl = doc.add_table(rows=6, cols=3)
    stack_tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    stk_headers = ["Category", "Technology / Library", "รายละเอียดและเหตุผลในการเลือกใช้ (Rationale)"]
    stk_hdr_cells = stack_tbl.rows[0].cells
    for i, title in enumerate(stk_headers):
        stk_hdr_cells[i].text = title
        set_cell_background(stk_hdr_cells[i], "0071E3")
        format_p(stk_hdr_cells[i].paragraphs[0], font_size=15, bold=True, color=COLOR_WHITE, align=WD_ALIGN_PARAGRAPH.CENTER)

    stack_data = [
        ("Frontend Web App Framework", "Next.js 16.2 (Turbopack App Router)\nReact 19 & TypeScript 5", "ประสิทธิภาพระดับสูงสุด โหลดหน้าจอเร็วทันใจ SSR/CSR Hybrid พร้อมระบบตรวจทานชนิดข้อมูล (Type-Safety)"),
        ("Styling & Responsive UI Design", "Vanilla CSS & Modern TailwindCSS\nCustom Design System Tokens", "ดีไซน์หรูหรา โทนสีน้ำเงินพรีเมียม (#0071e3) ตอบสนองทุกหน้าจอ (Responsive Mobile/Tablet/Desktop)"),
        ("Database & ORM Layer", "Prisma ORM 7.8 & Better-SQLite3 /\nPostgreSQL Multi-Database Adapter", "จัดการฐานข้อมูลรวดเร็ว ปลอดภัยด้วย Prepared Statements ป้องกัน SQL Injection 100%"),
        ("File & Media Processing Engine", "Sharp 0.35 & Adm-Zip & Excel XLSX Engine", "จัดการย่อขนาดรูปถ่ายความละเอียดสูง จัดการไฟล์ zip และนำเข้าแคตตาล็อกราคากลางจาก Excel"),
        ("Server & Process Management", "PM2 Enterprise Process Manager\nNode.js v20 LTS Server Environment", "ควบคุมการรันแอปพลิเคชันบน VPS Production (103.76.181.143) พร้อมระบบ Auto-Restart เมื่อเกิด Error"),
    ]

    for row_idx, data in enumerate(stack_data, start=1):
        row_cells = stack_tbl.rows[row_idx].cells
        if row_idx % 2 == 1:
            for c in row_cells: set_cell_background(c, "F9FAFB")
        row_cells[0].text = data[0]
        row_cells[1].text = data[1]
        row_cells[2].text = data[2]
        format_p(row_cells[0].paragraphs[0], font_size=14, bold=True, color=COLOR_DARK_BLUE)
        format_p(row_cells[1].paragraphs[0], font_size=14, color=COLOR_PRIMARY)
        format_p(row_cells[2].paragraphs[0], font_size=14, color=COLOR_TEXT)

    doc.add_page_break()

    # ==========================================
    # 6. PRICING PROPOSAL
    # ==========================================
    add_styled_heading("หมวดที่ 4: ข้อเสนอราคาและโมเดลการคิดค่าบริการ (Pricing Proposal)", level=1)

    add_styled_heading("4.1 โมเดลการคิดค่าบริการแบบ Pay-Per-Completed-Transaction", level=2)
    add_body_p("บริษัท เทคทันใจ อินโนเวชั่น จำกัด นำเสนอโมเดลการคิดค่าบริการที่เป็นธรรมและคุ้มค่าที่สุดสำหรับองค์กร โดยคิดค่าบริการเฉพาะ **'เคสใบเสนอราคาที่ถูกสแกน คุมราคา และอนุมัติเสร็จสิ้นสมบูรณ์ (Completed Transaction)'** เท่านั้น:")
    add_bullet_p("หากเคสถูกยกเลิก ไม่อนุมัติ หรือทดลองสแกนแล้วไม่ได้ใช้งานจริง จะไม่มีการคิดค่าบริการใดๆ ทั้งสิ้น (Zero Risk)", bold_prefix="• เสี่ยงเป็นศูนย์ (Zero Financial Risk): ")
    add_bullet_p("รวมค่าใช้งานเซิร์ฟเวอร์ Cloud VPS, ค่าบริการ AI API Engine, ระบบส่งอีเมล และการดูแลรักษาฟรี 24/7", bold_prefix="• รวมค่าบริการทุกอย่างแล้ว (All-Inclusive): ")

    add_styled_heading("4.2 ตารางอัตราค่าบริการตามปริมาณการใช้งาน (Volume Tiering Table)", level=2)

    pricing_tbl = doc.add_table(rows=5, cols=4)
    pricing_tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    p_headers = ["Package Tier", "ปริมาณเคสต่อเดือน (Volume / Month)", "ราคาต่อเคสที่อนุมัติสำเร็จ (Per Completed Transaction)", "เหมาะสำหรับ (Target Enterprise)"]
    p_hdr_cells = pricing_tbl.rows[0].cells
    for i, title in enumerate(p_headers):
        p_hdr_cells[i].text = title
        set_cell_background(p_hdr_cells[i], "0071E3")
        format_p(p_hdr_cells[i].paragraphs[0], font_size=14, bold=True, color=COLOR_WHITE, align=WD_ALIGN_PARAGRAPH.CENTER)

    pricing_data = [
        ("Starter Tier", "1 - 500 เคส / เดือน", "45 บาท / เคสที่อนุมัติสำเร็จ", "ศูนย์บริการ / อู่ซ่อมขนาดกลาง"),
        ("Standard Tier", "501 - 2,000 เคส / เดือน", "35 บาท / เคสที่อนุมัติสำเร็จ", "กลุ่มเครือข่ายอู่ซ่อม / บริษัทประกันขนาดกลาง"),
        ("Enterprise Tier", "2,001 - 5,000 เคส / เดือน", "25 บาท / เคสที่อนุมัติสำเร็จ", "บริษัทประกันภัยขนาดใหญ่ / TPA"),
        ("Volume Ultimate", "5,000+ เคสขึ้นไป / เดือน", "18 บาท / เคสที่อนุมัติสำเร็จ", "องค์กรประกันภัยชั้นนำระดับประเทศ"),
    ]

    for row_idx, data in enumerate(pricing_data, start=1):
        row_cells = pricing_tbl.rows[row_idx].cells
        if row_idx % 2 == 1:
            for c in row_cells: set_cell_background(c, "F9FAFB")
        row_cells[0].text = data[0]
        row_cells[1].text = data[1]
        row_cells[2].text = data[2]
        row_cells[3].text = data[3]
        format_p(row_cells[0].paragraphs[0], font_size=14, bold=True, color=COLOR_DARK_BLUE)
        format_p(row_cells[1].paragraphs[0], font_size=14, color=COLOR_TEXT)
        format_p(row_cells[2].paragraphs[0], font_size=14, bold=True, color=COLOR_PRIMARY)
        format_p(row_cells[3].paragraphs[0], font_size=14, color=COLOR_TEXT)

    doc.add_paragraph() # Spacer

    add_styled_heading("4.3 ประมาณการผลตอบแทนและการประหยัดต้นทุน (ROI Savings Analysis)", level=2)
    add_body_p("ตัวอย่างการคำนวณผลตอบแทนการลงทุนสำหรับบริษัทประกันภัยหรือเครือข่ายอู่ซ่อมที่มีปริมาณ 1,000 เคสต่อเดือน:")

    roi_tbl = doc.add_table(rows=5, cols=2)
    roi_tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    r_hdr_cells = roi_tbl.rows[0].cells
    r_hdr_cells[0].text = "รายการประเมินผลประโยชน์ (ROI Metric)"
    r_hdr_cells[1].text = "มูลค่าประเมินเปรียบเทียบ (Estimated Value)"
    set_cell_background(r_hdr_cells[0], "0F4C81")
    set_cell_background(r_hdr_cells[1], "0F4C81")
    format_p(r_hdr_cells[0].paragraphs[0], font_size=15, bold=True, color=COLOR_WHITE)
    format_p(r_hdr_cells[1].paragraphs[0], font_size=15, bold=True, color=COLOR_WHITE, align=WD_ALIGN_PARAGRAPH.RIGHT)

    roi_data = [
        ("ยอดการเสนอราคาเคลมรวม (1,000 เคส x เฉลี่ย 15,000 บาท/เคส)", "15,000,000 บาท / เดือน"),
        ("ยอดประหยัดจากการคุมราคากลางด้วย AI (ประหยัดเฉลี่ย 14%)", "2,100,000 บาท / เดือน"),
        ("ค่าบริการระบบ Claim-Thunjai AI (1,000 เคส x 35 บาท)", "35,000 บาท / เดือน"),
        ("ยอดประหยัดสุทธิคงเหลือหลังหักค่าบริการ (Net Savings ROI)", "2,065,000 บาท / เดือน (คุ้มค่า 59 เท่า)"),
    ]

    for row_idx, data in enumerate(roi_data, start=1):
        row_cells = roi_tbl.rows[row_idx].cells
        if row_idx == 4:
            set_cell_background(row_cells[0], "F0F7FF")
            set_cell_background(row_cells[1], "F0F7FF")
        elif row_idx % 2 == 1:
            set_cell_background(row_cells[0], "F9FAFB")
            set_cell_background(row_cells[1], "F9FAFB")

        row_cells[0].text = data[0]
        row_cells[1].text = data[1]
        format_p(row_cells[0].paragraphs[0], font_size=15, bold=(row_idx==4), color=COLOR_DARK_BLUE if row_idx!=4 else COLOR_PRIMARY)
        format_p(row_cells[1].paragraphs[0], font_size=15, bold=True, color=COLOR_PRIMARY if row_idx!=4 else COLOR_DARK_BLUE, align=WD_ALIGN_PARAGRAPH.RIGHT)

    doc.add_paragraph() # Spacer

    add_callout_box([
        "บริษัท เทคทันใจ อินโนเวชั่น จำกัด มุ่งมั่นที่จะมอบนวัตกรรมระบบคุมราคาเคลมที่ดีที่สุด ทรงประสิทธิภาพที่สุด และคุ้มค่าที่สุดสำหรับองค์กรของคุณ",
        "หากต้องการทดลองใช้งานระบบจริง (Live Demonstration) หรือสอบถามข้อมูลเพิ่มเติม สามารถติดต่อทีมงานได้ตลอด 24 ชั่วโมง",
        "\n• โทรศัพท์: 065-882-8333 | Email: athaporn@techthunjai.com | Line Official: @techthunjai",
        "• เว็บไซต์ระบบจริง: https://demo-claim.techthunjai.com/"
    ], title="บทสรุปและความพร้อมในการส่งมอบโครงการ")

    # Ensure data directory exists
    os.makedirs("data", exist_ok=True)
    file_path = os.path.abspath("data/Proposal_Claim_Thunjai_CEO.docx")
    doc.save(file_path)
    print("SUCCESS: Proposal docx generated at", file_path)

if __name__ == "__main__":
    create_proposal_docx()
